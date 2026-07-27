from __future__ import annotations

import io

import fitz
import pytest
from docx import Document
from starlette.datastructures import Headers, UploadFile

from app.core.config import Settings
from app.core.errors import AppError
from app.scenarios.recruitment.file_parser import ResumeFileParser


def upload(filename: str, content: bytes, content_type: str) -> UploadFile:
    return UploadFile(
        file=io.BytesIO(content),
        filename=filename,
        headers=Headers({"content-type": content_type}),
    )


def pdf_bytes(text: str | None) -> bytes:
    document = fitz.open()
    page = document.new_page()
    if text:
        page.insert_text((72, 72), text)
    content = document.tobytes()
    document.close()
    return content


@pytest.mark.asyncio
async def test_extracts_text_and_raw_file_metadata_from_pdf():
    parser = ResumeFileParser(Settings())
    source = pdf_bytes("Alice Example - Python FastAPI MySQL Docker backend engineer")

    parsed = await parser.parse(upload("candidate.pdf", source, "application/pdf"))

    assert "Alice Example" in parsed.text
    assert parsed.source_type == "pdf"
    assert parsed.source_size == len(source)
    assert len(parsed.source_sha256) == 64


@pytest.mark.asyncio
async def test_extracts_docx_paragraphs_and_tables():
    parser = ResumeFileParser(Settings())
    document = Document()
    document.add_paragraph("Alice Example has substantial backend engineering experience.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Python"
    output = io.BytesIO()
    document.save(output)

    parsed = await parser.parse(
        upload(
            "candidate.docx",
            output.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    )

    assert "backend engineering" in parsed.text
    assert "Skill | Python" in parsed.text


@pytest.mark.asyncio
async def test_rejects_scanned_pdf_without_text():
    parser = ResumeFileParser(Settings())

    with pytest.raises(AppError) as raised:
        await parser.parse(upload("scan.pdf", pdf_bytes(None), "application/pdf"))

    assert raised.value.code.value == "AI_RESUME_TEXT_NOT_FOUND"
    assert raised.value.http_status == 422


@pytest.mark.asyncio
async def test_rejects_file_over_configured_limit():
    parser = ResumeFileParser(Settings(recruitment_max_upload_mb=1))
    oversized = b"%PDF-" + b"x" * (1024 * 1024)

    with pytest.raises(AppError) as raised:
        await parser.parse(upload("large.pdf", oversized, "application/pdf"))

    assert raised.value.code.value == "AI_FILE_TOO_LARGE"
    assert raised.value.http_status == 413


@pytest.mark.asyncio
async def test_rejects_invalid_docx_signature():
    parser = ResumeFileParser(Settings())

    with pytest.raises(AppError) as raised:
        await parser.parse(
            upload(
                "candidate.docx",
                b"not a real docx but longer than twenty characters",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        )

    assert raised.value.code.value == "AI_FILE_CORRUPTED"
