from __future__ import annotations

import io
import json

from app.infrastructure.llm.catalog import ModelCatalogResult, get_model_catalog_client
from app.infrastructure.llm.dependencies import get_llm_client
from app.infrastructure.llm.models import LlmResponse, TokenUsage, UpstreamAttempt
from app.main import app


RESUME_TEXT = (
    "姓名：张三。学校：武汉理工大学。专业：软件工程。"
    "项目：使用 Spring Boot、Milvus、BM25 和大模型实现校园知识库问答。"
)
JOB_TEXT = (
    "招聘 AI 应用开发工程师，要求熟悉 Python 或 Java，具备后端服务开发能力，"
    "并有 RAG、大模型应用、Docker 或 Kubernetes 经验。"
)


def test_health_and_safe_settings(client):
    health = client.get("/api/v1/system/health")
    assert health.status_code == 200
    assert health.json()["data"]["status"] == "ok"
    assert health.headers["x-request-id"]

    response = client.get("/api/v1/settings")
    assert response.status_code == 200
    payload = response.json()["data"]
    assert "apiKeyConfigured" in payload
    assert "apiKey" not in payload
    assert "gptsapiApiKey" not in payload
    assert payload["adminAuthConfigured"] is True
    assert payload["speechMaxInputChars"] == 4096
    assert payload["speechMaxStreamChars"] == 50000


def test_admin_login_runtime_configuration_and_audit(client):
    unauthorized = client.put(
        "/api/v1/settings/llm",
        json={
            "baseUrl": "https://api.gptsapi.net/v1",
            "model": "test-model-b",
            "speechModel": "tts-1",
        },
    )
    assert unauthorized.status_code == 401

    invalid_login = client.post(
        "/api/v1/admin/login",
        json={"username": "test-admin", "password": "wrong-password"},
    )
    assert invalid_login.status_code == 401

    login = client.post(
        "/api/v1/admin/login",
        json={"username": "test-admin", "password": "test-admin-password"},
    )
    assert login.status_code == 200
    token = login.json()["data"]["accessToken"]
    admin_headers = {"Authorization": f"Bearer {token}"}

    class FakeCatalogClient:
        async def list_models(self, **_):
            return ModelCatalogResult(
                models=["test-model-a", "test-model-b", "tts-1", "tts-1-hd"],
                attempt=UpstreamAttempt(
                    attempt_no=1,
                    attempt_type="model_discovery",
                    status="success",
                    http_status=200,
                    error_code=None,
                    retryable=False,
                    duration_ms=3,
                ),
            )

    app.dependency_overrides[get_model_catalog_client] = lambda: FakeCatalogClient()
    try:
        models = client.get(
            "/api/v1/settings/models",
            params={"baseUrl": "https://api.gptsapi.net/v1"},
            headers=admin_headers,
        )
        assert models.status_code == 200
        assert models.json()["data"]["chatModels"] == ["test-model-a", "test-model-b"]
        assert models.json()["data"]["speechModels"] == ["tts-1", "tts-1-hd"]

        update = client.put(
            "/api/v1/settings/llm",
            json={
                "baseUrl": "https://api.gptsapi.net/v1",
                "model": "test-model-b",
                "speechModel": "tts-1-hd",
            },
            headers=admin_headers,
        )
        assert update.status_code == 200
        assert update.json()["data"]["model"] == "test-model-b"
        assert update.json()["data"]["speechModel"] == "tts-1-hd"
        assert update.json()["data"]["configurationSource"] == "database"

        rejected_url = client.put(
            "/api/v1/settings/llm",
            json={
                "baseUrl": "https://untrusted.example/v1",
                "model": "test-model-b",
                "speechModel": "tts-1",
            },
            headers=admin_headers,
        )
        assert rejected_url.status_code == 422

        audits = client.get(
            "/api/v1/settings/audits?page=1&pageSize=10",
            headers=admin_headers,
        )
        assert audits.status_code == 200
        items = audits.json()["data"]["items"]
        assert any(
            item["action"] == "settings.llm.update" and item["status"] == "success"
            for item in items
        )
        assert any(
            item["action"] == "settings.llm.update" and item["status"] == "failed"
            for item in items
        )
    finally:
        app.dependency_overrides.pop(get_model_catalog_client, None)

    logout = client.delete("/api/v1/admin/session", headers=admin_headers)
    assert logout.status_code == 200
    expired = client.get("/api/v1/admin/session", headers=admin_headers)
    assert expired.status_code == 401


def test_recruitment_flow_creates_audit_records(client):
    parse_response = client.post(
        "/api/v1/recruitment/resumes/parse",
        json={"resumeText": RESUME_TEXT},
        headers={"X-Caller-System": "pytest"},
    )
    assert parse_response.status_code == 200
    assert parse_response.json()["data"]["name"] == "张三"

    screening_response = client.post(
        "/api/v1/recruitment/screenings/evaluate",
        json={"resumeText": RESUME_TEXT, "jobDescription": JOB_TEXT},
        headers={"X-Caller-System": "pytest"},
    )
    assert screening_response.status_code == 200
    assert screening_response.json()["data"]["matchScore"] == 84

    interview_response = client.post(
        "/api/v1/recruitment/interview-kits/generate",
        json={
            "resumeText": RESUME_TEXT,
            "jobDescription": JOB_TEXT,
            "screeningRisks": ["评估口径待确认"],
        },
        headers={"X-Caller-System": "pytest"},
    )
    assert interview_response.status_code == 200
    assert interview_response.json()["data"]["questions"]

    audits = client.get("/api/v1/audits?page=1&pageSize=20")
    assert audits.status_code == 200
    assert audits.json()["data"]["total"] >= 3
    assert audits.json()["data"]["items"][0]["callerSystem"] == "pytest"
    assert audits.json()["data"]["items"][0]["createdAt"].endswith(("Z", "+00:00"))

    dashboard = client.get("/api/v1/dashboard/overview")
    assert dashboard.status_code == 200
    assert dashboard.json()["data"]["stats"]["businessRequests"] >= 3
    assert dashboard.json()["data"]["recentRequests"][0]["createdAt"].endswith(("Z", "+00:00"))
    usage_trend = dashboard.json()["data"]["usageTrend"]
    assert len(usage_trend) == 7
    assert sum(point["requestCount"] for point in usage_trend) >= 3


def test_tts_synthesis_returns_audio_and_creates_audit(client):
    response = client.post(
        "/api/v1/tts/synthesize",
        json={
            "text": "欢迎使用文字转语音助手。",
            "voice": "alloy",
            "responseFormat": "mp3",
            "speed": 1,
        },
        headers={"X-Caller-System": "pytest-tts"},
    )
    assert response.status_code == 200
    assert response.headers["content-type"].startswith("audio/wav")
    assert response.content.startswith(b"RIFF")
    assert response.headers["x-audio-model"]
    assert response.headers["x-audio-speed"] == "1"

    request_id = response.headers["x-request-id"]
    audits = client.get(f"/api/v1/audits?requestId={request_id}")
    assert audits.status_code == 200
    item = audits.json()["data"]["items"][0]
    assert item["businessCode"] == "tts"
    assert item["capabilityCode"] == "tts.speech.synthesize"
    assert item["requestMode"] == "binary"
    assert item["requestContentLength"] == len("欢迎使用文字转语音助手。")
    assert item["totalTokens"] == 0


def test_tts_stream_splits_long_text_and_creates_one_business_audit(client):
    text = ("这是用于验证长文本自动分段和流式播放的测试句子。" * 190).strip()
    assert len(text) > 4096

    response = client.post(
        "/api/v1/tts/synthesize-stream",
        json={
            "text": text,
            "voice": "alloy",
            "responseFormat": "mp3",
            "speed": 1,
        },
        headers={"X-Caller-System": "pytest-tts-stream"},
    )

    assert response.status_code == 200
    assert response.content.startswith(b"RIFF")
    assert response.headers["x-audio-streaming"] == "true"
    assert response.headers["x-audio-speed"] == "1"
    assert int(response.headers["x-audio-segments"]) >= 2

    request_id = response.headers["x-request-id"]
    audits = client.get(f"/api/v1/audits?requestId={request_id}")
    item = audits.json()["data"]["items"][0]
    assert item["businessCode"] == "tts"
    assert item["capabilityCode"] == "tts.speech.synthesize"
    assert item["requestMode"] == "stream"
    assert item["requestContentLength"] == len(text)
    assert item["status"] == "success"


def test_tts_stream_rejects_wav(client):
    response = client.post(
        "/api/v1/tts/synthesize-stream",
        json={
            "text": "流式播放只支持 MP3。",
            "voice": "alloy",
            "responseFormat": "wav",
            "speed": 1,
        },
    )

    assert response.status_code == 422
    assert response.json()["error"]["code"] == "AI_INVALID_REQUEST"


def test_validation_error_uses_standard_envelope(client):
    response = client.post("/api/v1/recruitment/resumes/parse", json={"resumeText": "太短"})
    assert response.status_code == 422
    payload = response.json()
    assert payload["success"] is False
    assert payload["error"]["code"] == "AI_INVALID_REQUEST"


def test_docx_resume_upload_creates_file_interface_audit(client):
    from docx import Document

    document = Document()
    document.add_heading("Candidate Resume", level=1)
    document.add_paragraph(
        "Alice Example, backend engineer with Python, FastAPI, MySQL and Docker experience."
    )
    buffer = io.BytesIO()
    document.save(buffer)

    response = client.post(
        "/api/v1/recruitment/resumes/parse-file",
        files={
            "file": (
                "candidate.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        headers={"X-Caller-System": "pytest-file"},
    )
    assert response.status_code == 200
    request_id = response.json()["requestId"]

    audits = client.get(f"/api/v1/audits?requestId={request_id}")
    assert audits.status_code == 200
    item = audits.json()["data"]["items"][0]
    assert item["interfacePath"] == "/api/v1/recruitment/resumes/parse-file"
    assert item["callerSystem"] == "pytest-file"
    assert item["upstreamCallCount"] == 0
    assert item["totalTokens"] == 0


def test_resume_upload_rejects_unsupported_type(client):
    response = client.post(
        "/api/v1/recruitment/resumes/parse-file",
        files={"file": ("candidate.txt", b"plain text resume content", "text/plain")},
    )
    assert response.status_code == 415
    assert response.json()["error"]["code"] == "AI_UNSUPPORTED_FILE_TYPE"


def test_format_repair_is_a_separate_upstream_call_not_a_retry(client):
    class RepairingClient:
        async def chat(self, request, attempt_type="primary"):
            usage = TokenUsage(prompt_tokens=2, completion_tokens=1, total_tokens=3)
            attempt = UpstreamAttempt(
                attempt_no=1,
                attempt_type=attempt_type,
                status="success",
                http_status=200,
                error_code=None,
                retryable=False,
                duration_ms=1,
                usage=usage,
            )
            if attempt_type == "primary":
                content = "not-json"
            else:
                assert RESUME_TEXT in request.messages[-1].content
                content = json.dumps(
                    {
                        "name": "Alice",
                        "school": None,
                        "major": None,
                        "graduationTime": None,
                        "skills": ["Python"],
                        "projects": [],
                    }
                )
            return LlmResponse(
                content=content,
                model=request.model,
                usage=usage,
                attempts=[attempt],
            )

    app.dependency_overrides[get_llm_client] = lambda: RepairingClient()
    try:
        response = client.post(
            "/api/v1/recruitment/resumes/parse",
            json={"resumeText": RESUME_TEXT},
            headers={"X-Caller-System": "pytest-repair"},
        )
    finally:
        app.dependency_overrides.pop(get_llm_client, None)

    assert response.status_code == 200
    request_id = response.json()["requestId"]
    audits = client.get(f"/api/v1/audits?requestId={request_id}")
    item = audits.json()["data"]["items"][0]
    assert item["upstreamCallCount"] == 2
    assert item["retryCount"] == 0
    assert item["totalTokens"] == 6


def test_empty_resume_result_triggers_format_repair(client):
    class EmptyThenRepairingClient:
        async def chat(self, request, attempt_type="primary"):
            usage = TokenUsage(prompt_tokens=2, completion_tokens=1, total_tokens=3)
            attempt = UpstreamAttempt(
                attempt_no=1,
                attempt_type=attempt_type,
                status="success",
                http_status=200,
                error_code=None,
                retryable=False,
                duration_ms=1,
                usage=usage,
            )
            if attempt_type == "primary":
                content = json.dumps(
                    {
                        "name": None,
                        "school": None,
                        "major": None,
                        "graduationTime": None,
                        "skills": [],
                        "projects": [],
                    }
                )
            else:
                assert RESUME_TEXT in request.messages[-1].content
                content = json.dumps(
                    {
                        "name": "Alice",
                        "school": None,
                        "major": None,
                        "graduationTime": None,
                        "skills": ["Python"],
                        "projects": [],
                    }
                )
            return LlmResponse(
                content=content,
                model=request.model,
                usage=usage,
                attempts=[attempt],
            )

    app.dependency_overrides[get_llm_client] = lambda: EmptyThenRepairingClient()
    try:
        response = client.post(
            "/api/v1/recruitment/resumes/parse",
            json={"resumeText": RESUME_TEXT},
            headers={"X-Caller-System": "pytest-empty-repair"},
        )
    finally:
        app.dependency_overrides.pop(get_llm_client, None)

    assert response.status_code == 200
    assert response.json()["data"]["name"] == "Alice"
    request_id = response.json()["requestId"]
    audits = client.get(f"/api/v1/audits?requestId={request_id}")
    item = audits.json()["data"]["items"][0]
    assert item["upstreamCallCount"] == 2
    assert item["retryCount"] == 0
    assert item["totalTokens"] == 6
