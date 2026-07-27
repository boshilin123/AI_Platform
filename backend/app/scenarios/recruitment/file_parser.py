from __future__ import annotations

import asyncio
import hashlib
import io
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path

import fitz
from docx import Document
from fastapi import UploadFile

from app.core.config import Settings
from app.core.error_codes import ErrorCode
from app.core.errors import AppError


PDF_MIME_TYPES = {"application/pdf", "application/octet-stream", ""}
DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream",
    "",
}
READ_CHUNK_SIZE = 1024 * 1024


@dataclass(frozen=True, slots=True)
class ParsedResumeInput:
    text: str
    source_type: str
    source_size: int
    source_sha256: str
    truncated: bool = False


class ResumeFileParser:
    """Validate and extract transient resume text without persisting source files."""

    def __init__(self, settings: Settings) -> None:
        self.max_upload_bytes = settings.recruitment_max_upload_mb * 1024 * 1024
        self.max_pdf_pages = settings.recruitment_max_pdf_pages
        self.max_extracted_chars = settings.recruitment_max_extracted_chars
        self.max_docx_uncompressed_bytes = (
            settings.recruitment_max_docx_uncompressed_mb * 1024 * 1024
        )

    async def parse(self, upload: UploadFile) -> ParsedResumeInput:
        source_type = self._source_type(upload.filename)
        self._validate_mime(source_type, upload.content_type)
        content = await self._read_limited(upload)
        source_sha256 = hashlib.sha256(content).hexdigest()
        text = await asyncio.to_thread(self._extract_text, source_type, content)
        normalized = self._normalize_text(text)
        if len(normalized) < 20:
            raise AppError(
                code=ErrorCode.RESUME_TEXT_NOT_FOUND,
                message="简历文件中未提取到足够的可读文本，暂不支持扫描版 PDF",
                http_status=422,
                retryable=False,
            )
        truncated = len(normalized) > self.max_extracted_chars
        if truncated:
            normalized = normalized[: self.max_extracted_chars]
        return ParsedResumeInput(
            text=normalized,
            source_type=source_type,
            source_size=len(content),
            source_sha256=source_sha256,
            truncated=truncated,
        )

    async def _read_limited(self, upload: UploadFile) -> bytes:
        chunks: list[bytes] = []
        total = 0
        while chunk := await upload.read(READ_CHUNK_SIZE):
            total += len(chunk)
            if total > self.max_upload_bytes:
                raise AppError(
                    code=ErrorCode.FILE_TOO_LARGE,
                    message=f"简历文件不能超过 {self.max_upload_bytes // 1024 // 1024}MB",
                    http_status=413,
                    retryable=False,
                )
            chunks.append(chunk)
        if not chunks:
            raise AppError(
                code=ErrorCode.FILE_CORRUPTED,
                message="上传的简历文件为空",
                http_status=422,
                retryable=False,
            )
        return b"".join(chunks)

    @staticmethod
    def _source_type(filename: str | None) -> str:
        suffix = Path(filename or "").suffix.lower()
        if suffix == ".pdf":
            return "pdf"
        if suffix == ".docx":
            return "docx"
        raise AppError(
            code=ErrorCode.UNSUPPORTED_FILE_TYPE,
            message="仅支持 PDF 和 DOCX 简历文件",
            http_status=415,
            retryable=False,
        )

    @staticmethod
    def _validate_mime(source_type: str, content_type: str | None) -> None:
        normalized = (content_type or "").lower()
        allowed = PDF_MIME_TYPES if source_type == "pdf" else DOCX_MIME_TYPES
        if normalized not in allowed:
            raise AppError(
                code=ErrorCode.UNSUPPORTED_FILE_TYPE,
                message="文件扩展名与 MIME 类型不匹配",
                http_status=415,
                retryable=False,
            )

    def _extract_text(self, source_type: str, content: bytes) -> str:
        if source_type == "pdf":
            return self._extract_pdf(content)
        return self._extract_docx(content)

    def _extract_pdf(self, content: bytes) -> str:
        if not content.startswith(b"%PDF-"):
            raise self._corrupted("PDF 文件签名无效")
        try:
            document = fitz.open(stream=content, filetype="pdf")
        except Exception as exc:
            raise self._corrupted("PDF 文件无法读取") from exc
        try:
            if document.needs_pass:
                raise AppError(
                    code=ErrorCode.PDF_ENCRYPTED,
                    message="暂不支持加密 PDF 简历",
                    http_status=422,
                    retryable=False,
                )
            if document.page_count > self.max_pdf_pages:
                raise AppError(
                    code=ErrorCode.FILE_TOO_LARGE,
                    message=f"PDF 简历不能超过 {self.max_pdf_pages} 页",
                    http_status=413,
                    retryable=False,
                )
            return "\n".join(page.get_text("text") for page in document)
        except AppError:
            raise
        except Exception as exc:
            raise self._corrupted("PDF 文本提取失败") from exc
        finally:
            document.close()

    def _extract_docx(self, content: bytes) -> str:
        if not content.startswith(b"PK"):
            raise self._corrupted("DOCX 文件签名无效")
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                names = set(archive.namelist())
                if "word/document.xml" not in names or "word/vbaProject.bin" in names:
                    raise self._corrupted("DOCX 文件结构无效")
                uncompressed_size = sum(item.file_size for item in archive.infolist())
                if uncompressed_size > self.max_docx_uncompressed_bytes:
                    raise AppError(
                        code=ErrorCode.FILE_TOO_LARGE,
                        message="DOCX 解压后内容超过安全限制",
                        http_status=413,
                        retryable=False,
                    )
            document = Document(io.BytesIO(content))
        except AppError:
            raise
        except (ValueError, KeyError, zipfile.BadZipFile) as exc:
            raise self._corrupted("DOCX 文件无法读取") from exc
        except Exception as exc:
            raise self._corrupted("DOCX 文本提取失败") from exc

        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    @staticmethod
    def _normalize_text(text: str) -> str:
        without_controls = "".join(
            char for char in text if char in "\n\t" or ord(char) >= 32
        )
        lines = [
            re.sub(r"[ \t]+", " ", line).strip()
            for line in without_controls.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        ]
        return "\n".join(line for line in lines if line).strip()

    @staticmethod
    def _corrupted(message: str) -> AppError:
        return AppError(
            code=ErrorCode.FILE_CORRUPTED,
            message=message,
            http_status=422,
            retryable=False,
        )
